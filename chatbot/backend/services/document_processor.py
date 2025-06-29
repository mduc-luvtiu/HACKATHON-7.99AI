import os
import PyPDF2
from docx import Document
from PIL import Image
import aiofiles
from typing import Optional, Dict, Any
import re
import google.generativeai as genai
import traceback
import base64
import io

class DocumentProcessor:
    def __init__(self, gemini_model=None):
        self.gemini_model = gemini_model
        
        # Use Gemini Vision model for image processing
        if genai:
            try:
                # Try different vision models
                try:
                    self.vision_model = genai.GenerativeModel('gemini-2.0-flash-exp')
                except:
                    try:
                        self.vision_model = genai.GenerativeModel('gemini-1.5-flash')
                    except:
                        self.vision_model = genai.GenerativeModel('gemini-1.5-pro')
                print(f"[DOC] Vision model initialized: {self.vision_model.model_name}")
            except Exception as e:
                print(f"[DOC] Failed to initialize vision model: {e}")
                self.vision_model = None
        else:
            self.vision_model = None
        
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image
        }
    
    async def process_document(self, file_path: str) -> str:
        print(f"[DOC][PROCESS] file_path={file_path}")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            processor = self.supported_extensions[file_extension]
            content = await processor(file_path)
            cleaned_content = self._clean_content(content)
            print(f"[DOC][CLEANED] {cleaned_content[:200]}")
            
            # Generate enhanced analysis if Gemini is available
            if self.gemini_model and cleaned_content:
                enhanced_analysis = await self._enhance_with_llm(cleaned_content, file_path)
                print(f"[DOC][LLM ENHANCED] {enhanced_analysis[:200]}")
                return enhanced_analysis
            
            return cleaned_content
        except Exception as e:
            print("[ERROR][DOC][PROCESS]", str(e))
            traceback.print_exc()
            raise Exception(f"Error processing document {file_path}: {str(e)}")

    async def _enhance_with_llm(self, content: str, file_path: str) -> str:
        """Enhance document content with LLM analysis"""
        try:
            # Create a comprehensive prompt for document analysis
            prompt = f"""
Phân tích tài liệu sau và cung cấp:

1. **Tóm tắt ngắn gọn** (2-3 câu)
2. **Các chủ đề chính** (3-5 chủ đề)
3. **Từ khóa quan trọng** (5-10 từ)
4. **Loại tài liệu** (báo cáo, hợp đồng, sách, v.v.)
5. **Độ tin cậy** (cao/trung bình/thấp)

Nội dung tài liệu:
{content[:2000]}...

Trả về kết quả dưới dạng JSON:
{{
    "summary": "Tóm tắt ngắn gọn",
    "topics": ["chủ đề 1", "chủ đề 2"],
    "keywords": ["từ khóa 1", "từ khóa 2"],
    "document_type": "loại tài liệu",
    "reliability": "cao/trung bình/thấp",
    "full_content": "Nội dung đầy đủ của tài liệu"
}}

Chỉ trả về JSON, không có text khác.
"""

            print(f"[DOC][LLM PROMPT] {prompt[:200]}")
            response = self.gemini_model.generate_content(prompt)
            response_text = response.text.strip()
            print(f"[DOC][LLM RESPONSE] {response_text[:200]}")
            
            # Try to parse JSON response
            import json
            result = json.loads(response_text)
            
            # Format the enhanced content
            enhanced_content = f"""
📄 **PHÂN TÍCH TÀI LIỆU**

📋 **Tóm tắt:** {result.get('summary', 'Không thể tóm tắt')}

🏷️ **Chủ đề chính:** {', '.join(result.get('topics', []))}

🔑 **Từ khóa:** {', '.join(result.get('keywords', []))}

📁 **Loại tài liệu:** {result.get('document_type', 'Không xác định')}

✅ **Độ tin cậy:** {result.get('reliability', 'Trung bình')}

---

📝 **NỘI DUNG ĐẦY ĐỦ:**
{result.get('full_content', content)}
"""
            
            return enhanced_content
            
        except Exception as e:
            print("[ERROR][DOC][LLM ENHANCE]", str(e))
            traceback.print_exc()
            return content
    
    async def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            content = ""
            async with aiofiles.open(file_path, 'rb') as file:
                pdf_content = await file.read()
                
            pdf_reader = PyPDF2.PdfReader(pdf_content)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                content += page.extract_text() + "\n"
            
            return content
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    async def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            return content
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    async def _process_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            return content
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    content = await file.read()
                return content
            except Exception as e:
                raise Exception(f"Error reading text file: {str(e)}")
    
    async def _process_image(self, file_path: str) -> str:
        """Extract text from image using Gemini Vision, always output in Vietnamese"""
        try:
            if not self.vision_model:
                print("[DOC] No vision model available, trying with regular Gemini model")
                if self.gemini_model:
                    try:
                        async with aiofiles.open(file_path, 'rb') as file:
                            image_data = await file.read()
                        image = Image.open(io.BytesIO(image_data))
                        if image.mode != 'RGB':
                            image = image.convert('RGB')
                        buffer = io.BytesIO()
                        image.save(buffer, format='JPEG')
                        image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                        prompt = """
Hãy đọc và trích xuất toàn bộ văn bản có trong hình ảnh này. Nếu văn bản là tiếng Anh, hãy dịch toàn bộ sang tiếng Việt. Chỉ trả về kết quả cuối cùng bằng tiếng Việt, giữ nguyên cấu trúc và ý nghĩa. Nếu không có văn bản, trả về 'Không có văn bản trong hình ảnh.'
"""
                        response = self.gemini_model.generate_content([
                            prompt,
                            {
                                "mime_type": "image/jpeg",
                                "data": image_base64
                            }
                        ])
                        extracted_text = response.text.strip()
                        if not extracted_text or extracted_text.lower() in ["không có văn bản", "no text found", "empty"]:
                            return "Không có văn bản trong hình ảnh."
                        return extracted_text
                    except Exception as e:
                        print(f"[DOC] Regular Gemini model failed: {e}")
                        return f"Không thể xử lý ảnh với Gemini model: {str(e)}"
                else:
                    return "Không thể xử lý ảnh vì chưa có Gemini model."
            print(f"[DOC] Using vision model: {self.vision_model.model_name}")
            async with aiofiles.open(file_path, 'rb') as file:
                image_data = await file.read()
            image = Image.open(io.BytesIO(image_data))
            if image.mode != 'RGB':
                image = image.convert('RGB')
            buffer = io.BytesIO()
            image.save(buffer, format='JPEG')
            image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
            prompt = """
Hãy đọc và trích xuất toàn bộ văn bản có trong hình ảnh này. Nếu văn bản là tiếng Anh, hãy dịch toàn bộ sang tiếng Việt. Chỉ trả về kết quả cuối cùng bằng tiếng Việt, giữ nguyên cấu trúc và ý nghĩa. Nếu không có văn bản, trả về 'Không có văn bản trong hình ảnh.'
"""
            response = self.vision_model.generate_content([
                prompt,
                {
                    "mime_type": "image/jpeg",
                    "data": image_base64
                }
            ])
            extracted_text = response.text.strip()
            if not extracted_text or extracted_text.lower() in ["không có văn bản", "no text found", "empty"]:
                return "Không có văn bản trong hình ảnh."
            return extracted_text
        except Exception as e:
            print(f"[ERROR][IMAGE PROCESSING] {str(e)}")
            traceback.print_exc()
            return f"Lỗi khi xử lý ảnh: {str(e)}"
    
    def _clean_content(self, content: str) -> str:
        """Clean and normalize extracted content"""
        if not content:
            return ""
        
        # Remove extra whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Remove special characters but keep Vietnamese diacritics
        content = re.sub(r'[^\w\sàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ.,!?;:()[\]{}"\'-]', ' ', content)
        
        # Remove multiple spaces
        content = re.sub(r'\s+', ' ', content)
        
        # Remove leading/trailing whitespace
        content = content.strip()
        
        return content
    
    async def get_document_summary(self, content: str, max_length: int = 500) -> str:
        print(f"[DOC][SUMMARY][INPUT] {content[:200]}")
        if not content:
            return "No content available"
        if self.gemini_model:
            try:
                prompt = f"""
Tóm tắt đoạn văn bản sau bằng tiếng Việt trong {max_length} ký tự, giữ lại thông tin quan trọng nhất. Nếu văn bản là tiếng Anh, hãy dịch tóm tắt sang tiếng Việt:

{content}

Tóm tắt:
"""
                print(f"[DOC][SUMMARY][PROMPT] {prompt[:200]}")
                response = self.gemini_model.generate_content(prompt)
                print(f"[DOC][SUMMARY][RESPONSE] {response.text[:200]}")
                return response.text.strip()
            except Exception as e:
                print("[ERROR][DOC][SUMMARY]", str(e))
                traceback.print_exc()
                return self._simple_summary(content, max_length)
        else:
            return self._simple_summary(content, max_length)
    
    def _simple_summary(self, content: str, max_length: int = 500) -> str:
        """Simple summary without LLM"""
        # Simple summary: take first few sentences
        sentences = re.split(r'[.!?]+', content)
        summary = ""
        
        for sentence in sentences:
            if len(summary + sentence) < max_length:
                summary += sentence.strip() + ". "
            else:
                break
        
        return summary.strip() if summary else content[:max_length] + "..."
    
    async def extract_keywords(self, content: str, max_keywords: int = 10) -> list:
        print(f"[DOC][KEYWORDS][INPUT] {content[:200]}")
        if not content:
            return []
        
        if self.gemini_model:
            try:
                prompt = f"""
Trích xuất {max_keywords} từ khóa quan trọng nhất từ đoạn văn bản sau. Trả về dưới dạng JSON array:

{content[:1000]}...

Format: ["từ khóa 1", "từ khóa 2", ...]
"""
                print(f"[DOC][KEYWORDS][PROMPT] {prompt[:200]}")
                response = self.gemini_model.generate_content(prompt)
                response_text = response.text.strip()
                print(f"[DOC][KEYWORDS][RESPONSE] {response_text[:200]}")
                
                # Try to parse JSON response
                import json
                keywords = json.loads(response_text)
                return keywords[:max_keywords]
            except Exception as e:
                print("[ERROR][DOC][KEYWORDS]", str(e))
                traceback.print_exc()
                return self._simple_keyword_extraction(content, max_keywords)
        else:
            return self._simple_keyword_extraction(content, max_keywords)
    
    def _simple_keyword_extraction(self, content: str, max_keywords: int = 10) -> list:
        """Simple keyword extraction without LLM"""
        # Simple keyword extraction based on frequency
        words = re.findall(r'\b\w+\b', content.lower())
        
        # Filter out common words
        stop_words = {
            'và', 'của', 'là', 'có', 'được', 'trong', 'với', 'từ', 'này', 'đó',
            'the', 'and', 'of', 'to', 'a', 'in', 'is', 'it', 'you', 'that'
        }
        
        word_freq = {}
        for word in words:
            if len(word) > 2 and word not in stop_words:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Sort by frequency and return top keywords
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:max_keywords]]

    async def answer_question_about_document(self, content: str, question: str) -> str:
        if not self.gemini_model:
            return "Tôi không thể trả lời câu hỏi về tài liệu này vì chưa có LLM."
        try:
            prompt = f"""
Dựa trên nội dung tài liệu sau, hãy trả lời câu hỏi một cách chính xác và trung thực. Nếu không có thông tin trong tài liệu, hãy nói rõ.

Nội dung tài liệu:
{content}

Câu hỏi: {question}

Trả lời:
"""
            response = self.gemini_model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            return f"Xin lỗi, tôi không thể trả lời câu hỏi này: {str(e)}"

# Khi upload tài liệu thành công, lưu nội dung vào chat_service.document_memory[user_id] = [list các đoạn] 